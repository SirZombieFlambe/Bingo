import random
import platform

import docx
from docx.shared import Inches, Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from tkinter import Tk, filedialog
import time

def get_random_bingo_section(bingo_answers, is_file=True):
    if is_file:
        try:
            with open(bingo_answers, "r", encoding="utf-8") as file:
                answer = file.readlines()
            file.close()
            return random.choice(answer)
        except FileNotFoundError:
            print(f"There is no{bingo_answers} file. Please try again")

    else:
        return random.choice(bingo_answers)


def make_doc_table(table_array, docs_number, doc_group, title):
    table_info = table_array.copy()

    document = docx.Document()

    heading = document.add_heading(title, level=0)
    heading.alignment = 1

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)

    rows = 5
    cols = 5
    table_info.insert(12, 'FREE')

    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # print(table_info)
    j = 0
    print(table_info)
    for row in range(rows):
        row_cells = table.rows[row].cells
        for col in range(cols):
            cell = row_cells[col]
            cell.text = table_info[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Inches(1.5)
            cell.height = Inches(1)
            j = j + 1

    for row in table.rows:
        row.height = Inches(1.5)

    document.save(f'bingo_card_{doc_group}{docs_number}.docx')


def main():
    doc_group = 'a'

    repeat = True
    while repeat:

        title = input("What is the bingo game's title? Ex Lawrence Bingo: ")
        valid_number = False
        docs_number = 0
        while not valid_number:
            docs_number = input("How many Bingo Cards do you want: ")
            if docs_number.isdigit():
                valid_number = True
            else:
                print("Hey that's not a number. Try again.")

        allowed_duplicated = False

        duplicates = input("Would you like to allow for duplicates, if not ensure that there are enough elements (Y/n): ")

        if duplicates.upper() == 'Y':
            allowed_duplicated = True
        else:
            print("Duplicates are not allowed. This might cause an error")

        user_settings = input('Press "m" to enter bingo elements manually, or press any other key to select a .txt '
                              'file: ')

        card_size_max = 25

        if user_settings == 'm':
            bingo_card_elements = []
            if platform.system() == "Windows":
                eof_key = "<Ctrl+Z>"
            else:
                eof_key = "<Ctrl+D>"

            input_mode = True

            while input_mode:
                try:
                    user_input = input(f"Enter next value or {eof_key} to stop: ")
                except EOFError:  # Ctrl+Z + Enter on Windows, Ctrl+D on Linux/MacOSX
                    input_mode = False


                except KeyboardInterrupt:  # Ctrl + C - will exit program immediately if not caught
                    break
                else:
                    bingo_card_elements.append(user_input)



            for num in range(int(docs_number)):

                bingo_card = []
                card_size = 1
                full_card = False
                while not full_card:

                    current_element = get_random_bingo_section(bingo_card_elements, is_file=False)

                    if current_element not in bingo_card:
                        bingo_card.append(current_element)
                        card_size = card_size + 1
                        if card_size == card_size_max:
                            full_card = True
                    elif allowed_duplicated:
                        bingo_card.append(current_element)
                        card_size = card_size + 1
                        if card_size == card_size_max:
                            full_card = True

                make_doc_table(bingo_card, num + 1, doc_group, title)

        else:

            # Use tkinter to open Windows Explorer and select a file

            root = Tk()
            root.withdraw()  # Hide the main window
            file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
            root.destroy()  # Close the main window after file selection

            if not file_path:
                print("No file selected. Exiting.")
                return

            for num in range(int(docs_number)):
                bingo_card = []
                full_card = False
                card_size = 1
                while not full_card:
                    current_element = get_random_bingo_section(file_path).rsplit('\n')[0]

                    if current_element not in bingo_card:

                        bingo_card.append(current_element)

                        card_size = card_size + 1
                        if card_size == card_size_max:
                            full_card = True

                    elif allowed_duplicated:
                        bingo_card.append(current_element)

                        card_size = card_size + 1
                        if card_size == card_size_max:
                            full_card = True


                make_doc_table(bingo_card, num + 1, doc_group, title)

        repeat_query = input("Would you like to make more cards Y/n, or press any key to end: ")
        if repeat_query.upper() == 'Y':
            repeat = True
            doc_group = chr(ord(doc_group) + 1)
            print("\n\n")
        else:
            repeat = False
            print("\n\n\nHave a nice day :)")
            time.sleep(3)



if __name__ == '__main__':
    main()
